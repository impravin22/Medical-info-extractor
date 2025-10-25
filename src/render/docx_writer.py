from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document

from src.domain.layout import LayoutSpec
from src.domain.schema import Event, Report


def _format_date(d: date) -> str:
    return d.strftime("%m/%d/%Y")


def _event_cell(event: Event, column: str) -> str:
    if column == "Date":
        return _format_date(event.date)
    if column == "Provider":
        return event.provider
    if column == "Physician":
        return event.physician or ""
    if column == "Facility":
        return event.facility or ""
    if column == "Reason":
        return event.reason_raw
    if column == "Summary":
        return event.reason_summary or event.reason_raw
    if column == "Ref":
        return str(event.ref_page)
    return ""


def render_report_to_docx(
    report: Report,
    out_path: str | Path,
    layout: LayoutSpec,
    template_path: Optional[str | Path] = None,
) -> None:
    doc = Document(str(template_path)) if template_path else Document()

    # Header block
    h = report.claimant
    doc.add_heading("MEDICAL SUMMARY", level=1)

    # Line 1: RE, SSN, Title, DLI - use extracted data only
    title_part = f"Title: {h.title}" if h.title else ""
    dli_part = f"DLI: {_format_date(h.dli)}" if h.dli else ""
    line1_parts = [f"RE: {h.name}", f"SSN: {h.ssn}", title_part, dli_part]
    doc.add_paragraph("  ".join(p for p in line1_parts if p))

    # Line 2: AOD, DOB, Ages - compute from extracted data only
    age_at_aod = ""
    current_age = ""

    if h.aod and h.dob:
        age_at_aod = str(
            (h.aod.year - h.dob.year)
            - (1 if (h.aod.month, h.aod.day) < (h.dob.month, h.dob.day) else 0)
        )

    if h.dob:
        today = datetime.today().date()
        current_age = str(
            (today.year - h.dob.year)
            - (1 if (today.month, today.day) < (h.dob.month, h.dob.day) else 0)
        )

    aod_str = _format_date(h.aod) if h.aod else ""
    dob_str = _format_date(h.dob) if h.dob else ""
    doc.add_paragraph(
        f"AOD: {aod_str}  Date of Birth: {dob_str}  Age at AOD: {age_at_aod}  Current Age: {current_age}"
    )

    # Line 3: Education - use extracted data only
    if h.last_grade_completed or h.attended_special_ed is not None:
        grade = h.last_grade_completed or ""
        special_ed = ""
        if h.attended_special_ed is True:
            special_ed = "Yes"
        elif h.attended_special_ed is False:
            special_ed = "No"

        doc.add_paragraph(
            f"Last Grade Completed: {grade}  Attended Special Ed Classes: {special_ed}"
        )

    if report.alleged_impairments:
        doc.add_paragraph("Alleged impairments:")
        for imp in report.alleged_impairments:
            doc.add_paragraph(f"• {imp}")

    # Last Updated line
    doc.add_paragraph(f"Last Updated: {datetime.today().strftime('%A, %B %d, %Y')}")

    # Timeline table
    # If the template already contains a table, reuse the first one; else create
    table = (
        doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=len(layout.columns))
    )
    # Ensure header row exists and has correct number of columns
    while len(table.rows[0].cells) < len(layout.columns):
        table.add_column(100000)
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(layout.columns):
        hdr_cells[idx].text = col

    # Apply sorting
    events = list(report.events)
    for key, direction in reversed(layout.sort_by):
        reverse = direction == "desc"
        if key == "Date":
            events.sort(key=lambda e: e.date, reverse=reverse)
        elif key == "Provider":
            events.sort(key=lambda e: e.provider.lower(), reverse=reverse)
        elif key == "Ref":
            events.sort(key=lambda e: e.ref_page, reverse=reverse)

    # Rows
    for event in events:
        row_cells = table.add_row().cells
        for c_idx, col in enumerate(layout.columns):
            val = _event_cell(event, col)
            if col == "Ref" and val.isdigit():
                val = f"Pg {val}"
            row_cells[c_idx].text = val

    # Note missing columns
    if layout.unknown_columns:
        doc.add_paragraph(
            "Unavailable columns: " + ", ".join(sorted(set(layout.unknown_columns)))
        )

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
